from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "图形化版贪吃蛇概要设计说明书_新版.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 38, 48)
MUTED = RGBColor(95, 108, 118)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
TABLE_WIDTH_DXA = 9360


def set_run(run, size=11, color=INK, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("图形化贪吃蛇概要设计")
    set_run(run, size=9.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("BUPT EasyX Edition")
    set_run(run, size=9, color=MUTED)


def paragraph(doc, text="", size=11, color=INK, bold=False, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run(run, size=size, color=color, bold=bold)
    return p


def heading(doc, text, level=1):
    doc.add_paragraph(text, style=f"Heading {level}")


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run(run, size=11)


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run(run, size=11)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in [("top", 80), ("bottom", 80), ("start", 120), ("end", 120)]:
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    set_table_width(t, widths)
    set_cell_margins(t)

    for idx, header in enumerate(headers):
        cell = t.rows[0].cells[idx]
        shade(cell, HEADER_FILL)
        set_cell_width(cell, widths[idx])
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run(run, size=10.5, bold=True)

    for values in rows:
        row = t.add_row()
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_run(run, size=10.2)

    paragraph(doc, "", after=2)
    return t


def callout(doc, title, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    set_table_width(t, [TABLE_WIDTH_DXA])
    set_cell_margins(t)
    cell = t.rows[0].cells[0]
    set_cell_width(cell, TABLE_WIDTH_DXA)
    shade(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run(run, size=11, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_run(run, size=10.5)
    paragraph(doc, "", after=2)


def build_doc():
    doc = Document()
    setup_document(doc)

    paragraph(doc, "概要设计说明书", size=11, color=MUTED, bold=True, after=4)
    paragraph(doc, "图形化版贪吃蛇游戏", size=24, bold=True, after=4)
    paragraph(
        doc,
        "Visual Studio + EasyX 图形化实现，覆盖大地图、分辨率设置、AI 对战道具、音效与模块化扩展接口。",
        size=12.5,
        color=MUTED,
        after=10,
    )
    table(
        doc,
        ["项目", "内容"],
        [
            ("工程名称", "BUPT_Eating_Snake_Graphics"),
            ("开发环境", "Visual Studio 2022，EasyX，Win32 输入函数，WinMM 音频接口"),
            ("核心文件", "include/*.h，src/*.c，assets/*，tools/*.py"),
            ("生成日期", "2026 年 6 月 3 日"),
        ],
        [2200, 7160],
    )

    callout(
        doc,
        "版本目标",
        "本版本在 OJ 版贪吃蛇规则基础上完成图形化重构，支持单人模式、AI 对战模式、限时挑战模式、皮肤切换、设置界面、音效系统和可扩展道具系统。"
    )

    heading(doc, "1 需求与实现范围")
    table(
        doc,
        ["需求", "实现方案"],
        [
            ("EasyX 贴图显示", "Render_loadSkin 按皮肤目录加载 PNG；格子尺寸变化时重新按当前 cellSize 加载，使用 3 参数 putimage 绘制。"),
            ("欢迎界面", "菜单包含单人模式、AI 对战模式、限时挑战模式、更换时装、设置、退出游戏。"),
            ("Win32 控制", "input.c 使用 GetAsyncKeyState 检测 W/A/S/D、E、1、2、方向键、Enter、Esc 等按键。"),
            ("模块化管理", "规则、AI、渲染、输入、音频、UI 流程分别放入独立源文件，通过头文件暴露接口。"),
            ("AI 对战增强", "AI 对战只生成食物、弓箭、护盾、尖刺、减速时钟；AI 根据难度采取不同道具获取和攻击策略。"),
            ("限时挑战", "默认 90 秒，玩家在时间内争取更高分，支持多样模式下的食物、陷阱、护盾等道具。"),
            ("多地图与分辨率", "设置支持 20x20、50x50、100x100；窗口支持紧凑、标准、大三档分辨率，大窗口下 100x100 可达到 10px 单格。"),
            ("开局等待", "进入游戏后先绘制地图和提示，检测到 W/A/S/D 后再移动和计时。"),
            ("速度调节", "游戏中按 1 加速、2 减速；速度档位影响移动间隔，并按比例改变食物奖励分数。"),
        ],
        [2500, 6860],
    )

    heading(doc, "2 总体架构")
    paragraph(
        doc,
        "工程采用分层结构。规则层不直接绘图或播放声音，只维护 GameState；UI 层读取输入、调用规则更新、消费声音事件；渲染层只根据 GameState 画当前画面。",
    )
    table(
        doc,
        ["模块", "文件", "职责"],
        [
            ("公共定义", "include/common.h", "常量、枚举、Pos、Snake、GameConfig、GameState、SoundEvent。"),
            ("游戏规则", "src/game.c", "地图初始化、道具生成、移动、增长、碰撞、胜负、速度档位、弓箭结算。"),
            ("AI 决策", "src/ai.c", "低/中/高难度寻路、空间评估、道具价值评估和攻击性评分。"),
            ("输入", "src/input.c", "Win32 键盘检测和菜单按键去抖。"),
            ("渲染", "src/render_easyx.c", "EasyX 窗口、菜单、设置页、地图、蛇、贴图、动态棋盘尺寸。"),
            ("音频", "src/audio.c", "背景音乐循环播放和事件音效播放。"),
            ("界面流程", "src/ui.c", "欢迎界面、设置界面、模式选择、开局等待、游戏循环、结算界面。"),
            ("入口", "src/main.c", "初始化模块、保存设置、启动不同模式、释放资源。"),
        ],
        [1500, 2300, 5560],
    )

    heading(doc, "3 数据结构设计")
    table(
        doc,
        ["结构/枚举", "关键字段", "说明"],
        [
            ("GameConfig", "mode、variant、aiDifficulty、resolution、mapSize、enableStepGrowth、musicEnabled、soundEnabled", "保存一局游戏启动前的配置，设置界面直接修改这些字段。"),
            ("GameState", "cells、player、ai、speedLevel、remainingSeconds、soundEvents", "保存运行时地图、两条蛇、速度档位、计时和待播放音效。"),
            ("Snake", "body、length、dir、nextDir、score、slowMs、bowArrows、shieldCharges", "body[0] 是蛇头；AI 对战资源保存在 bowArrows 和 shieldCharges。"),
            ("CellType", "CELL_FOOD、CELL_OBSTACLE、CELL_BATTLE_BOW、CELL_BATTLE_SPIKE 等", "地图格子的静态内容，蛇身单独保存在 Snake.body 中。"),
            ("SoundEvent", "SOUND_EAT_NORMAL、SOUND_BOW、SOUND_ARROW_HIT 等", "规则层写入事件，UI 层调用 Audio_playEvent 播放。"),
        ],
        [1700, 3900, 4760],
    )

    heading(doc, "4 地图与大尺寸适配")
    bullet(doc, "Game_validMapSize 将地图限制在 20、50、100 三档，所有遍历只访问当前 mapSize 范围。")
    bullet(doc, "障碍物数量按 mapSize 面积缩放，避免 50x50 和 100x100 地图过空。")
    bullet(doc, "普通模式的食物、奖励食物、陷阱和护盾数量按地图大小提高；AI 对战模式使用独立的新道具生成表。")
    bullet(doc, "RenderContext 保存 boardPixelSize，cellSize = boardPixelSize / mapSize；贴图按 cellSize 重载，避免大地图切换后显示错位。")
    bullet(doc, "大分辨率选项将棋盘提升到 1000 像素，100x100 地图下每格 10 像素，50x50 地图下每格 20 像素。")

    heading(doc, "5 设置界面设计")
    table(
        doc,
        ["设置项", "取值", "影响"],
        [
            ("N 步自动增长", "开启 / 关闭", "开启后每 DEFAULT_GROWTH_INTERVAL 步自动增长一次；关闭后只通过食物增长。"),
            ("地图尺寸", "20x20 / 50x50 / 100x100", "影响地图遍历、障碍物和道具数量、AI BFS 范围以及绘制格子大小。"),
            ("窗口分辨率", "1080x820 / 1440x1000 / 1700x1080", "影响 EasyX 窗口大小和棋盘像素大小。"),
            ("背景音乐", "开启 / 关闭", "控制 bgm.wav 是否循环播放。"),
            ("游戏音效", "开启 / 关闭", "控制吃食物、道具、碰撞、射箭等音效是否播放。"),
        ],
        [2100, 3300, 4960],
    )

    heading(doc, "6 AI 对战新道具")
    table(
        doc,
        ["道具", "玩家效果", "AI 处理策略"],
        [
            ("弓箭", "吃到后 bowArrows 加 1；按 E 沿当前蛇头方向发射。命中头部直接胜利，命中身体则删除命中段及其后方尾段。障碍物和墙壁会阻挡。", "低难度只在能射头时发射；中/高难度会射身体或头部，高难度还会主动争夺弓箭。"),
            ("护盾", "获得一次 shieldCharges，可抵挡一次弓箭或尖刺陷阱。", "AI 会在没有护盾时提高道具价值，且允许带盾踩尖刺。"),
            ("尖刺陷阱", "无护盾碰到直接死亡，有护盾则消耗护盾并清除陷阱。", "AI 寻路会把无护盾尖刺视为不可通行；有护盾时可通过但仍有评分惩罚。"),
            ("减速时钟", "吃到后让敌方 slowMs=1000，移动和方向变化都会变慢。", "高难度在距离玩家较近时更重视时钟，用于追击和压迫。"),
        ],
        [1600, 4860, 3900],
    )

    heading(doc, "7 AI 难度策略")
    table(
        doc,
        ["难度", "决策方式", "特点"],
        [
            ("低", "根据道具价值和曼哈顿距离贪心移动，发箭条件严格。", "容易预测，攻击性较低，主要用于练习。"),
            ("中", "用 BFS 找到目标道具，检查第一步后的可达空间，避免明显死路。", "更稳定，会利用弓箭、护盾、时钟，但不过度冒险。"),
            ("高", "对四个方向评分：可达空间、食物距离、玩家可走方向数、玩家距离、战斗道具价值、是否存在弓箭命中线。", "主动抢资源、压缩玩家路线，并合理使用弓箭和减速时钟。"),
        ],
        [1200, 5260, 3760],
    )

    heading(doc, "8 移动、碰撞与计分")
    number(doc, "每次更新先读取玩家方向，再根据移动间隔累积 moveTimerMs。")
    number(doc, "到达移动时刻后构造 StepPlan，判断墙、障碍、陷阱、尖刺、蛇身和对战碰撞。")
    number(doc, "如果本步增长，则蛇长加一；否则尾部自然覆盖。N 步增长和吃食物增长可以同时触发。")
    number(doc, "AI 对战中两条蛇同步计算下一格，若同格相撞或蛇头互换位置则判平局。")
    number(doc, "speedLevel 范围为 -2 到 2；正数缩短移动间隔并提高食物分数，负数延长移动间隔并降低食物分数。")

    heading(doc, "9 资源与扩展接口")
    table(
        doc,
        ["扩展目标", "修改位置", "说明"],
        [
            ("新增道具", "common.h、game.c、render_easyx.c、audio.c", "增加 CellType、生成数量、效果处理、贴图枚举和必要音效。"),
            ("新增皮肤", "assets/<skin>/、render_easyx.c", "放入同名 PNG，并在 SKINS 表增加目录和显示名。"),
            ("新增音效", "assets/audio/、common.h、audio.c", "放入 WAV，增加 SoundEvent，并在 soundPath 中映射。"),
            ("新增设置项", "GameConfig、Ui_runSettings、Render_drawSettings", "新增配置字段、菜单修改逻辑和显示行。"),
            ("新增 AI 策略", "ai.c", "调整 itemValueForAi、hardCandidateScore 或新增难度分支。"),
            ("新增关卡", "game.c", "在 Game_init 前后增加关卡加载函数，根据文件填充 cells。"),
        ],
        [1800, 3200, 4360],
    )

    heading(doc, "10 测试方案")
    bullet(doc, "编译检查：使用 Visual Studio x64 Debug/Release 构建，确认 EasyX 和 winmm.lib 链接正常。")
    bullet(doc, "地图测试：分别进入 20x20、50x50、100x100，观察棋盘尺寸、贴图比例、障碍和道具数量。")
    bullet(doc, "设置测试：切换 N 步增长、分辨率、音乐、音效，确认立即生效或在下一局生效。")
    bullet(doc, "开局测试：进入游戏后不按 W/A/S/D 时地图保持静止；按任意方向后开始移动。")
    bullet(doc, "AI 对战测试：分别测试低、中、高难度；验证弓箭、护盾、尖刺、减速时钟效果。")
    bullet(doc, "速度测试：按 1/2 改变速度档位，观察移动间隔与吃食物分数变化。")
    bullet(doc, "音效测试：开始、吃食物、吃弓箭、命中箭、护盾、陷阱、碰撞都能播放不同 WAV。")

    heading(doc, "11 已知约束")
    bullet(doc, "EasyX 环境需要在 Visual Studio 中安装配置；本地脚本生成资源不依赖 EasyX。")
    bullet(doc, "100x100 地图受屏幕高度限制，最大分辨率下单格约 10 像素；若需要更大格子，可继续增加滚动画布或摄像机视口。")
    bullet(doc, "当前弓箭为瞬时射线结算，没有绘制飞行动画；后续可在 GameState 中增加 Arrow 实体做动画。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
